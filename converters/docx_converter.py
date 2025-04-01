"""
使用python-docx和BeautifulSoup实现的HTML转DOCX转换器
支持中文论文格式要求、CSS样式解析和class属性应用
"""
import io
import re
import logging
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Length
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from converters.base_converter import BaseConverter

# 配置日志
logger = logging.getLogger(__name__)
# 如果已经有root logger配置了，就不再添加处理器
if not logger.parent.handlers and not logger.handlers:
    logger.setLevel(logging.DEBUG)
    handler = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)


class DocxConverter(BaseConverter):
    """使用python-docx和BeautifulSoup的转换器实现，支持中文论文格式要求"""
    
    # 中文字号到磅值的映射
    CN_FONT_SIZE_MAP = {
        '初号': Pt(42),
        '小初': Pt(36),
        '一号': Pt(26),
        '小一': Pt(24),
        '二号': Pt(22),
        '小二': Pt(18),
        '三号': Pt(16),
        '小三': Pt(15),
        '四号': Pt(14),
        '小四': Pt(12),
        '五号': Pt(10.5),
        '小五': Pt(9),
        '六号': Pt(7.5),
        '小六': Pt(6.5),
        '七号': Pt(5.5),
        '八号': Pt(5)
    }
    
    # CSS字体系列映射到Word字体
    FONT_FAMILY_MAP = {
        'SimSun': '宋体',
        'SimHei': '黑体',
        'Times New Roman': 'Times New Roman',
        'Arial': 'Arial'
    }
    
    def convert(self, html_content: str) -> bytes:
        """
        将HTML内容转换为DOCX格式，支持样式和class属性
        
        Args:
            html_content: HTML字符串内容
            
        Returns:
            DOCX文档的二进制数据
        """
        logger.info("开始转换HTML到DOCX")
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        
        # 提取样式信息
        css_styles = self._extract_css_styles(soup)
        
        # 处理文档主体
        self._process_html_body(soup, doc, css_styles)
        
        # 将文档保存到内存缓冲区
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        logger.info("HTML到DOCX转换完成")
        return docx_bytes.getvalue()
    
    def _extract_css_styles(self, soup):
        """
        从HTML中提取CSS样式
        
        Args:
            soup: BeautifulSoup对象
            
        Returns:
            样式字典，格式为 {selector: {property: value}}
        """
        styles = {}
        
        # 查找所有style标签
        style_tags = soup.find_all('style')
        logger.info(f"找到 {len(style_tags)} 个样式标签")
        
        for style_tag in style_tags:
            css_text = style_tag.string
            if not css_text:
                logger.warning("样式标签内容为空")
                continue
                
            logger.debug(f"原始CSS文本：\n{css_text}")
            
            # 先移除CSS注释
            css_text = re.sub(r'/\*.*?\*/', '', css_text, flags=re.DOTALL)
            logger.debug(f"移除注释后的CSS文本：\n{css_text}")
                
            # 解析CSS规则 - 使用更强大的正则表达式来处理多行规则
            rules = re.findall(r'([^{]+){([^}]+)}', css_text, re.DOTALL)
            logger.info(f"从CSS中提取到 {len(rules)} 条样式规则")
            
            for selector, properties in rules:
                # 移除选择器中可能残留的注释并清理空格
                selector = selector.strip()
                logger.debug(f"处理原始选择器: '{selector}'")
                
                # 跳过空选择器
                if not selector:
                    continue
                
                # 处理多个选择器（以逗号分隔）
                for single_selector in selector.split(','):
                    single_selector = single_selector.strip()
                    if not single_selector:
                        continue
                    
                    logger.debug(f"处理单个选择器: '{single_selector}'")
                    
                    # 解析属性
                    style_properties = {}
                    for prop in properties.split(';'):
                        if ':' not in prop:
                            continue
                        
                        name, value = prop.split(':', 1)
                        name = name.strip()
                        # 清除引号和转义字符
                        value = value.strip().replace('\\', '').replace('"', '').replace("'", '')
                        
                        if value:
                            style_properties[name] = value
                    
                    styles[single_selector] = style_properties
                    logger.debug(f"最终选择器及其样式: '{single_selector}' => {style_properties}")
        
        logger.info(f"CSS样式解析完成，共 {len(styles)} 个选择器")
        logger.debug(f"所有解析的选择器: {list(styles.keys())}")
        return styles
    
    def _selector_matches_element(self, selector, element):
        """
        判断选择器是否匹配元素
        
        Args:
            selector: CSS选择器字符串
            element: BeautifulSoup元素
            
        Returns:
            布尔值，表示是否匹配
        """
        # 处理元素类型选择器 (如 p, h1, h2 等)
        if selector == element.name:
            return True
            
        # 处理类选择器 (如 .abstract, .toc 等)
        if selector.startswith('.'):
            class_name = selector[1:]
            return class_name in element.get('class', [])
            
        # 处理元素类型+类选择器 (如 h1.chapter-title, p.english 等)
        if '.' in selector and not selector.startswith('.'):
            element_type, class_part = selector.split('.', 1)
            return element.name == element_type and class_part in element.get('class', [])
            
        # 处理组合选择器 (如 .abstract h1, .abstract .summary-title 等)
        if ' ' in selector:
            # 获取最后一个部分，它应该匹配当前元素
            parts = selector.split()
            last_part = parts[-1]
            
            # 检查当前元素是否匹配最后一部分
            if not self._selector_matches_element(last_part, element):
                return False
                
            # 检查祖先元素是否匹配前面的部分
            ancestor_parts = parts[:-1]
            current = element.parent
            while current and ancestor_parts:
                if self._selector_matches_element(ancestor_parts[-1], current):
                    ancestor_parts.pop()
                    if not ancestor_parts:
                        return True
                current = current.parent
                
            # 如果检查完所有祖先元素后还有未匹配的部分，则不匹配
            return len(ancestor_parts) == 0
            
        return False
        
    def _get_element_styles(self, element, css_styles):
        """
        获取元素的样式信息
        
        Args:
            element: BeautifulSoup元素
            css_styles: CSS样式字典
            
        Returns:
            元素的样式字典
        """
        element_styles = {}
        
        # 检查元素类名
        class_names = element.get('class', [])
        element_type = element.name
        
        logger.debug(f"处理元素: {element_type}, 类名: {class_names}")
        logger.debug(f"可用的CSS选择器: {list(css_styles.keys())[:10]}..." if len(css_styles) > 10 else f"可用的CSS选择器: {list(css_styles.keys())}")
        
        # 按CSS选择器优先级排序和应用样式
        selector_matches = []
        
        # 收集所有匹配的选择器
        for selector, properties in css_styles.items():
            if self._selector_matches_element(selector, element):
                # 计算选择器优先级 (简化版)
                specificity = 0
                if selector == element_type:  # 元素类型选择器，优先级最低
                    specificity = 1
                elif selector.startswith('.'):  # 类选择器
                    specificity = 10
                elif '.' in selector and not selector.startswith('.'):  # 元素类型+类选择器
                    specificity = 11
                elif ' ' in selector:  # 组合选择器
                    # 简单计算，每个部分都增加优先级
                    specificity = 100 + len(selector.split())
                    
                selector_matches.append((selector, properties, specificity))
                
        # 按优先级排序
        selector_matches.sort(key=lambda x: x[2])
        
        # 应用样式，后应用的会覆盖先应用的
        for selector, properties, specificity in selector_matches:
            element_styles.update(properties)
            logger.debug(f"应用选择器样式: {selector} (优先级: {specificity}) => {properties}")
        
        # 处理内联样式 (最高优先级)
        inline_style = element.get('style', '')
        if inline_style:
            for prop in inline_style.split(';'):
                if ':' not in prop:
                    continue
                name, value = prop.split(':', 1)
                element_styles[name.strip()] = value.strip()
            
            logger.debug(f"应用内联样式: {inline_style}")
        
        if element_styles:
            logger.debug(f"最终应用到元素 {element_type} 的样式: {element_styles}")
        else:
            logger.debug(f"元素 {element_type} 未匹配到任何样式")
            
        return element_styles
    
    def _apply_paragraph_styles(self, paragraph, styles):
        """
        应用段落样式
        
        Args:
            paragraph: python-docx段落对象
            styles: 样式字典
        """
        if not styles:
            return
            
        logger.debug(f"应用段落样式: {styles}")
        
        # 字体系列
        font_family = styles.get('font-family', '')
        if font_family:
            # 处理多个字体系列（以逗号分隔）
            font_families = [f.strip() for f in font_family.split(',')]
            
            # 使用第一个能识别的字体
            for font in font_families:
                font = font.strip('"\'')  # 移除引号
                if font in self.FONT_FAMILY_MAP:
                    font_name = self.FONT_FAMILY_MAP[font]
                    for run in paragraph.runs:
                        run.font.name = font_name
                        # 对中文字符应用中文字体
                        if font_name in ['宋体', '黑体']:
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    logger.debug(f"应用字体: {font} => {font_name}")
                    break
        
        # 字体大小
        font_size = styles.get('font-size', '')
        if font_size:
            # 处理中文字号格式 (如 "小四号")
            if font_size in self.CN_FONT_SIZE_MAP:
                pt_size = self.CN_FONT_SIZE_MAP[font_size]
                logger.debug(f"应用中文字号: {font_size} => {pt_size}")
            # 处理pt格式 (如 "12pt")
            else:
                size_match = re.match(r'(\d+\.?\d*)(?:pt|px)?', font_size)
                if size_match:
                    size_value = float(size_match.group(1))
                    # 如果是px单位，做简单转换 (1px ≈ 0.75pt)
                    if 'px' in font_size:
                        size_value *= 0.75
                    pt_size = Pt(size_value)
                    logger.debug(f"应用字体大小: {font_size} => {pt_size}")
                else:
                    pt_size = None
                    logger.warning(f"无法解析字体大小: {font_size}")
                
            if pt_size:
                for run in paragraph.runs:
                    run.font.size = pt_size
        
        # 字体颜色
        font_color = styles.get('color', '')
        if font_color:
            # 处理颜色代码
            try:
                # 尝试解析颜色值（支持#RRGGBB格式）
                if font_color.startswith('#'):
                    r, g, b = int(font_color[1:3], 16), int(font_color[3:5], 16), int(font_color[5:7], 16)
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(r, g, b)
                    logger.debug(f"应用字体颜色: {font_color} => RGB({r},{g},{b})")
            except Exception as e:
                logger.warning(f"无法解析字体颜色: {font_color}, 错误: {e}")
        else:
            # 设置默认黑色，覆盖Word模板中可能的其他颜色
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            logger.debug("应用默认黑色字体")
        
        # 字体粗细
        font_weight = styles.get('font-weight', '')
        if font_weight == 'bold':
            for run in paragraph.runs:
                run.font.bold = True
            logger.debug("应用粗体")
        
        # 字体样式 (斜体)
        font_style = styles.get('font-style', '')
        if font_style == 'italic':
            for run in paragraph.runs:
                run.font.italic = True
            logger.debug("应用斜体")
        
        # 对齐方式
        text_align = styles.get('text-align', '')
        if text_align == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            logger.debug("应用居中对齐")
        elif text_align == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            logger.debug("应用右对齐")
        elif text_align == 'justify':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            logger.debug("应用两端对齐")
        
        # 行距
        line_height = styles.get('line-height', '')
        if line_height:
            try:
                # 处理数值格式 (如 "1.5")
                line_height_value = float(line_height)
                paragraph.paragraph_format.line_spacing = line_height_value
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                logger.debug(f"应用行距: {line_height}")
            except ValueError:
                # 可能是其他格式，如 "normal"
                logger.warning(f"无法解析行距: {line_height}")
        
        # 段前距离
        margin_top = styles.get('margin-top', '')
        if margin_top:
            if margin_top.endswith('em'):
                try:
                    # 提取数值部分
                    margin_match = re.match(r'(\d+\.?\d*)em', margin_top)
                    if margin_match:
                        value = float(margin_match.group(1))
                        paragraph.paragraph_format.space_before = Pt(value * 12)  # 假设1em = 12pt
                        logger.debug(f"应用段前距离: {margin_top}")
                except ValueError:
                    logger.warning(f"无法解析段前距离: {margin_top}")
        
        # 首行缩进 - 使用CSS的text-indent属性或为普通段落添加默认缩进
        text_indent = styles.get('text-indent', '')
        if text_indent:
            try:
                # 处理em单位 (如 "2em")
                if text_indent.endswith('em'):
                    indent_match = re.match(r'(\d+\.?\d*)em', text_indent)
                    if indent_match:
                        value = float(indent_match.group(1))
                        paragraph.paragraph_format.first_line_indent = Pt(value * 12)  # 假设1em = 12pt
                        logger.debug(f"应用首行缩进: {text_indent}")
                # 处理px或pt单位 (如 "28px"或"21pt")
                elif text_indent.endswith('px') or text_indent.endswith('pt'):
                    indent_match = re.match(r'(\d+\.?\d*)(?:px|pt)', text_indent)
                    if indent_match:
                        value = float(indent_match.group(1))
                        # 如果是px单位，做简单转换 (1px ≈ 0.75pt)
                        if text_indent.endswith('px'):
                            value *= 0.75
                        paragraph.paragraph_format.first_line_indent = Pt(value)
                        logger.debug(f"应用首行缩进: {text_indent}")
                # 处理cm单位 (如 "0.5cm")
                elif text_indent.endswith('cm'):
                    indent_match = re.match(r'(\d+\.?\d*)cm', text_indent)
                    if indent_match:
                        value = float(indent_match.group(1))
                        # 1cm ≈ 28.3pt
                        paragraph.paragraph_format.first_line_indent = Pt(value * 28.3)
                        logger.debug(f"应用首行缩进: {text_indent}")
            except ValueError as e:
                logger.warning(f"无法解析首行缩进: {text_indent}, 错误: {e}")
        else:
            # 为普通段落添加默认缩进 (2个中文字符宽度，约28磅)
            # 简化判断逻辑，仅使用paragraph.style.name == 'Normal'
            if paragraph.style.name == 'Normal':
                # 检查段落是否已有首行缩进设置
                if paragraph.paragraph_format.first_line_indent is None:
                    paragraph.paragraph_format.first_line_indent = Pt(28)
                    logger.debug("应用默认首行缩进: 28pt (约2个中文字符)")
    
    def _process_html_body(self, soup, doc, css_styles):
        """
        处理HTML主体内容
        
        Args:
            soup: BeautifulSoup对象
            doc: python-docx文档对象
            css_styles: CSS样式字典
        """
        body = soup.body if soup.body else soup
        logger.info("开始处理HTML主体内容")
        
        # 应用body元素的样式到文档默认样式
        body_styles = self._get_element_styles(body, css_styles)
        if body_styles:
            # 应用body样式到文档样式
            default_style = doc.styles['Normal']
            
            # 字体系列
            font_family = body_styles.get('font-family', '')
            if font_family:
                font_families = [f.strip() for f in font_family.split(',')]
                for font in font_families:
                    font = font.strip('"\'')
                    if font in self.FONT_FAMILY_MAP:
                        font_name = self.FONT_FAMILY_MAP[font]
                        default_style.font.name = font_name
                        if font_name in ['宋体', '黑体']:
                            default_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        break
            
            # 字体大小
            font_size = body_styles.get('font-size', '')
            if font_size:
                if font_size in self.CN_FONT_SIZE_MAP:
                    default_style.font.size = self.CN_FONT_SIZE_MAP[font_size]
                else:
                    size_match = re.match(r'(\d+\.?\d*)(?:pt|px)?', font_size)
                    if size_match:
                        size_value = float(size_match.group(1))
                        if 'px' in font_size:
                            size_value *= 0.75
                        default_style.font.size = Pt(size_value)
            
            # 行距
            line_height = body_styles.get('line-height', '')
            if line_height:
                try:
                    line_height_value = float(line_height)
                    default_style.paragraph_format.line_spacing = line_height_value
                    default_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                except ValueError:
                    pass
        
        # 遍历所有元素
        self._process_element(body, doc, css_styles)
        
        logger.info("HTML主体内容处理完成")
    
    def _process_element(self, element, doc, css_styles, current_section=None):
        """
        递归处理HTML元素
        
        Args:
            element: BeautifulSoup元素
            doc: python-docx文档对象
            css_styles: CSS样式字典
            current_section: 当前处理的文档部分
        """
        # 如果是文本节点，直接返回
        if element.name is None:
            return
        
        # 获取元素样式
        element_styles = self._get_element_styles(element, css_styles)
        class_names = element.get('class', [])
        
        # 处理表格元素
        if element.name == 'table':
            logger.info("处理表格元素")
            self._process_table(element, doc, css_styles, current_section)
            return
            
        # 处理不同类型的元素
        if element.name == 'div':
            # 特殊处理不同部分
            if 'cover' in class_names:
                current_section = 'cover'
                logger.info("处理文档封面部分")
            elif 'toc' in class_names:
                current_section = 'toc'
                logger.info("处理目录部分")
            elif 'abstract' in class_names:
                current_section = 'abstract'
                logger.info("处理摘要部分")
            elif 'keywords' in class_names:
                current_section = 'keywords'
                logger.info("处理关键词部分")
            elif 'references' in class_names:
                current_section = 'references'
                logger.info("处理参考文献部分")
            elif 'acknowledgments' in class_names:
                current_section = 'acknowledgments'
                logger.info("处理致谢部分")
            
            # 处理子元素
            for child in element.children:
                self._process_element(child, doc, css_styles, current_section)
            # 已经处理了子元素，直接返回，避免重复处理
            return
                
        elif element.name in ['h1', 'h2', 'h3']:
            # 处理标题
            level = int(element.name[1])
            text = element.get_text().strip()
            logger.info(f"处理{level}级标题: {text}")
            
            # 添加标题段落 - 使用add_heading保持标题层级，但立即设置颜色
            if 'chapter-title' in class_names:
                # 一级(章)标题
                para = doc.add_heading(text, level=1)
                logger.debug("添加章标题")
            elif 'section-title' in class_names:
                # 二级(节)标题
                para = doc.add_heading(text, level=2)
                logger.debug("添加节标题")
            elif 'subsection-title' in class_names:
                # 三级标题
                para = doc.add_heading(text, level=3)
                logger.debug("添加小节标题")
            else:
                para = doc.add_heading(text, level=level)
                logger.debug(f"添加{level}级标题")
            
            # 立即设置标题颜色为黑色
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            logger.debug("设置标题颜色为黑色")
            
            # 特殊处理摘要标题（在abstract部分内的h1标签）
            if current_section == 'abstract' and element.name == 'h1':
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in para.runs:
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    run.font.size = Pt(16)
                logger.debug("设置摘要标题居中样式")
            
            # 应用CSS样式
            self._apply_paragraph_styles(para, element_styles)
            
        elif element.name == 'p':
            # 处理段落
            # 特殊处理关键词标题
            if current_section == 'keywords' and element.find('strong', class_='keyword-title'):
                # 处理包含关键词标题的段落，只加粗关键词标题部分
                keyword_title = element.find('strong', class_='keyword-title').get_text()
                keywords_content = element.get_text()[len(keyword_title):].strip()
                
                # 创建空段落
                para = doc.add_paragraph()
                
                # 添加加粗的关键词标题
                title_run = para.add_run(keyword_title)
                title_run.font.bold = True
                title_run.font.name = '宋体'
                title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                # 添加不加粗的关键词内容
                content_run = para.add_run(keywords_content)
                content_run.font.bold = False
                content_run.font.name = '宋体'
                content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                logger.debug(f"处理关键词段落：标题'{keyword_title}'已加粗，内容保持普通样式")
                
                # 应用段落样式
                self._apply_paragraph_styles(para, element_styles)
                
            else:
                # 处理普通段落
                text = element.get_text().strip()
                para = doc.add_paragraph(text)
                logger.debug(f"添加段落: {text[:30]}...")
                
                # 特殊处理摘要标题
                if current_section == 'abstract' and 'summary-title' in class_names:
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in para.runs:
                        run.font.name = '黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run.font.size = Pt(16)
                    logger.debug("设置摘要标题样式")
                
                # 处理英文段落
                if 'english' in class_names:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                    logger.debug("设置英文段落字体")
                
                # 应用段落样式
                self._apply_paragraph_styles(para, element_styles)
            
        elif element.name == 'ul' or element.name == 'ol':
            # 处理列表
            logger.info(f"处理{'无序' if element.name == 'ul' else '有序'}列表")
            for li in element.find_all('li', recursive=False):
                text = li.get_text().strip()
                para = doc.add_paragraph(text)
                if element.name == 'ul':
                    para.style = 'List Bullet'
                else:
                    para.style = 'List Number'
                
                logger.debug(f"添加列表项: {text[:30]}...")
                
                # 应用段落样式
                self._apply_paragraph_styles(para, element_styles)
                
        elif element.name == 'img' and element.get('src'):
            # 处理图片 (实际应用中可能需要进一步完善)
            src = element.get('src')
            logger.info(f"处理图片: {src}")
            try:
                doc.add_picture(src, width=Inches(6.0))
                logger.debug("图片添加成功")
            except Exception as e:
                # 添加图片失败时，添加替代文本
                alt = element.get('alt', src)
                para = doc.add_paragraph(f"[图片: {alt}]")
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                logger.warning(f"图片添加失败: {e}")
        
        # 递归处理子元素 (仅针对非div和未直接处理子元素的元素)
        if element.name not in ['div', 'p', 'h1', 'h2', 'h3', 'ul', 'ol']:
            for child in element.children:
                if child.name:  # 只处理元素节点，跳过文本节点
                    self._process_element(child, doc, css_styles, current_section) 
    
    def _process_table(self, element, doc, css_styles, current_section=None):
        """
        处理HTML表格元素，转换为DOCX表格
        
        Args:
            element: BeautifulSoup表格元素
            doc: python-docx文档对象
            css_styles: CSS样式字典
            current_section: 当前处理的文档部分
            
        Returns:
            None
        """
        # 分析表格结构
        rows = element.find_all('tr', recursive=True)
        if not rows:
            logger.warning("表格中未找到行元素")
            return
            
        # 确定表格最大列数
        max_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'], recursive=False)
            col_count = sum(int(cell.get('colspan', 1)) for cell in cells)
            max_cols = max(max_cols, col_count)
            
        if max_cols == 0:
            logger.warning("表格无有效列")
            return
            
        logger.info(f"创建表格: {len(rows)}行 x {max_cols}列")
        
        # 创建DOCX表格
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'  # 设置基本网格样式
        
        # 处理表格整体样式
        table_styles = self._get_element_styles(element, css_styles)
        self._apply_table_styles(table, table_styles)
        
        # 处理行和单元格
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'], recursive=False)
            col_index = 0
            
            for cell in cells:
                # 获取单元格跨行和跨列信息
                rowspan = int(cell.get('rowspan', 1))
                colspan = int(cell.get('colspan', 1))
                
                # 跳过已经被合并的单元格
                while col_index < max_cols and table.cell(i, col_index)._element.getparent() != table.rows[i]._element:
                    col_index += 1
                    
                if col_index >= max_cols:
                    break
                    
                # 获取目标单元格
                target_cell = table.cell(i, col_index)
                
                # 合并单元格
                if rowspan > 1 or colspan > 1:
                    # 确保不超出表格边界
                    end_row = min(i + rowspan - 1, len(rows) - 1)
                    end_col = min(col_index + colspan - 1, max_cols - 1)
                    
                    if end_row > i or end_col > col_index:
                        target_cell = table.cell(i, col_index)
                        merge_cell = table.cell(end_row, end_col)
                        target_cell.merge(merge_cell)
                        logger.debug(f"合并单元格: 从({i},{col_index})到({end_row},{end_col})")
                
                # 处理单元格内容
                cell_content = cell.get_text().strip()
                if cell_content:
                    # 清除现有段落内容
                    for paragraph in target_cell.paragraphs:
                        if paragraph.text:
                            for run in paragraph.runs:
                                run.text = ""
                    
                    # 添加文本
                    paragraph = target_cell.paragraphs[0] if target_cell.paragraphs else target_cell.add_paragraph()
                    run = paragraph.add_run(cell_content)
                    
                    # 判断是否表头单元格
                    if cell.name == 'th':
                        run.bold = True
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # 应用单元格样式
                    cell_styles = self._get_element_styles(cell, css_styles)
                    self._apply_cell_styles(target_cell, paragraph, run, cell_styles)
                    self._apply_paragraph_styles(paragraph, cell_styles)
                
                # 移动到下一个列索引
                col_index += colspan
        
        logger.info("表格处理完成")
    
    def _apply_table_styles(self, table, styles):
        """应用样式到整个表格"""
        if not styles:
            return
            
        # 表格宽度
        width = styles.get('width', '')
        if width:
            try:
                # 解析百分比或像素值
                if '%' in width:
                    # 百分比转换为近似英寸
                    percent = float(width.strip('%')) / 100
                    table.width = Inches(6 * percent)  # 假设页面宽度为6英寸
                else:
                    # 像素转换为英寸 (假设96dpi)
                    pixels = float(re.match(r'(\d+)', width).group(1))
                    table.width = Inches(pixels / 96)
            except (ValueError, AttributeError):
                logger.warning(f"无法解析表格宽度: {width}")
        
        # 边框样式
        border = styles.get('border', '')
        if border:
            # 简单处理边框样式
            try:
                for cell in table._cells:
                    self._set_cell_border(cell, border)
            except Exception as e:
                logger.warning(f"设置表格边框时出错: {e}")
    
    def _apply_cell_styles(self, cell, paragraph, run, styles):
        """应用样式到表格单元格"""
        if not styles:
            return
            
        # 背景色
        bg_color = styles.get('background-color', '')
        if bg_color:
            try:
                # 解析RGB颜色
                if bg_color.startswith('#'):
                    # 16进制颜色值
                    r = int(bg_color[1:3], 16)
                    g = int(bg_color[3:5], 16)
                    b = int(bg_color[5:7], 16)
                    self._set_cell_background(cell, r, g, b)
                elif bg_color.startswith('rgb'):
                    # RGB函数格式
                    rgb_match = re.match(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', bg_color)
                    if rgb_match:
                        r, g, b = map(int, rgb_match.groups())
                        self._set_cell_background(cell, r, g, b)
            except Exception as e:
                logger.warning(f"设置单元格背景色时出错: {e}")
        
        # 文本对齐
        text_align = styles.get('text-align', '')
        if text_align:
            if text_align == 'center':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif text_align == 'right':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif text_align == 'left':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif text_align == 'justify':
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    def _set_cell_border(self, cell, border_spec):
        """设置单元格边框"""
        # 简化实现：暂不支持详细的边框样式
        pass
    
    def _set_cell_background(self, cell, r, g, b):
        """设置单元格背景色"""
        try:
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            # 创建单元格阴影元素
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self._rgb_to_hex(r, g, b)}"/>')
            
            # 应用到单元格属性
            cell._tc.get_or_add_tcPr().append(shd)
        except ImportError:
            logger.warning("无法设置单元格背景色，需要安装lxml库")
    
    def _rgb_to_hex(self, r, g, b):
        """将RGB值转换为十六进制颜色代码"""
        return f"{r:02x}{g:02x}{b:02x}" 